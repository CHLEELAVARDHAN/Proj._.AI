from flask import (
    Flask, render_template, request, redirect, url_for,
    flash, session, jsonify, send_file, send_from_directory
)
import os
import re
import json
import io
import requests
import pandas as pd
from datetime import datetime
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from docx import Document

# optional collaboration module (your existing). Provide safe fallback if missing.
try:
    from modules.collaboration import collaborate, show_feed
except Exception:
    def collaborate(*args, **kwargs):
        return None
    def show_feed(*args, **kwargs):
        return []

# Load .env
load_dotenv()

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Folders & files
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
DATA_FOLDER = os.path.join(BASE_DIR, "data")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DATA_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

IDEAS_FILE = os.path.join(DATA_FOLDER, "ideas.json")
USERS_FILE = os.path.join(DATA_FOLDER, "users.json")
CSV_FILE = os.path.join(DATA_FOLDER, "skills_companies_packages.csv")

ALLOWED_EXTENSIONS = {"pdf", "docx"}

# Helper: safe filename for downloads
def _safe_filename(name: str) -> str:
    base = re.sub(r'[^A-Za-z0-9 _.-]+', '', (name or "")).strip()
    return (base or "project").replace(' ', '_')[:60]

# -------------------------
# Project Uploads History Helpers
# -------------------------
UPLOAD_HISTORY_FILE = os.path.join(DATA_FOLDER, "uploads_history.json")

def load_uploads():
    try:
        with open(UPLOAD_HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}
    except json.JSONDecodeError:
        return {}

def save_uploads(data):
    with open(UPLOAD_HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

# -------------------------
# Ideas persistence helpers
# ensures list structure, assigns incremental ids
# -------------------------
def load_ideas():
    try:
        with open(IDEAS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            # support legacy dict-wrapped structure
            if isinstance(data, dict) and "ideas" in data:
                data = data["ideas"]
            if isinstance(data, list):
                return data
            return []
    except FileNotFoundError:
        return []
    except json.JSONDecodeError:
        return []

def save_ideas(ideas):
    # ensure it's a list
    if not isinstance(ideas, list):
        raise ValueError("ideas must be a list")
    with open(IDEAS_FILE, "w", encoding="utf-8") as f:
        json.dump(ideas, f, indent=4, ensure_ascii=False)

def next_idea_id(ideas):
    # return next numeric id
    if not ideas:
        return 1
    try:
        ids = [int(i.get("id", 0)) for i in ideas]
        return max(ids) + 1
    except Exception:
        return len(ideas) + 1

# -------------------------
# Users persistence helpers
# -------------------------
def load_users():
    try:
        with open(USERS_FILE, "r", encoding="utf-8") as f:
            txt = f.read().strip()
            return json.loads(txt) if txt else {}
    except FileNotFoundError:
        return {}
    except json.JSONDecodeError:
        return {}

def save_users(users):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, indent=4, ensure_ascii=False)

# -------------------------
# Skills / companies load (CSV optional)
# -------------------------
skills_data = {}
if os.path.exists(CSV_FILE):
    try:
        df = pd.read_csv(CSV_FILE)
        for _, row in df.iterrows():
            if not {"skill", "company", "package"}.issubset(row.index):
                continue
            skill = str(row["skill"]).strip().lower()
            company = str(row["company"]).strip()
            package = str(row["package"]).strip()
            skills_data.setdefault(skill, []).append({"name": company, "package": package})
    except Exception:
        skills_data = {}

# fallback sample if CSV not present
if not skills_data:
    skills_data = {
        "python": [{"name": "Acme", "package": "5 LPA"}],
        "data science": [{"name": "DataCorp", "package": "6 LPA"}],
        "java": [{"name": "BigSoft", "package": "4 LPA"}]
    }

# -------------------------
# Utilities
# -------------------------
def encode_skill(skill: str) -> str:
    return skill.replace("/", "-").replace(" ", "_").replace("+", "plus").lower()

def decode_skill(encoded: str) -> str:
    return encoded.replace("-", "/").replace("_", " ").replace("plus", "+")

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# -------------------------
# Authentication Routes
# -------------------------
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        users = load_users()
        if username in users and check_password_hash(users[username]["password"], password):
            session["username"] = username
            flash(f"Welcome {username}!", "success")
            return redirect(url_for("welcome"))
        flash("Invalid username or password.", "danger")
    return render_template("login.html")

@app.route("/download/<filename>", endpoint="download_project_file")
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        confirm = request.form.get("confirm_password", "")
        if password != confirm:
            flash("Passwords do not match.", "danger")
            return render_template("register.html")
        users = load_users()
        if username in users:
            flash("Username already exists.", "warning")
            return render_template("register.html")
        users[username] = {"email": email, "password": generate_password_hash(password)}
        save_users(users)
        flash("Registration successful! Please login.", "success")
        return redirect(url_for("login"))
    return render_template("register.html")

from flask import send_file
from docx import Document
import io

@app.route('/start_project/<int:idea_id>')
def start_projects(idea_id):
    ideas = load_ideas()  # load ideas.json
    if 0 <= idea_id < len(ideas):
        idea = ideas[idea_id]
        
        # Create Word document in memory
        doc = Document()
        doc.add_heading(idea["idea"], level=1)
        doc.add_paragraph(f"Sector: {idea.get('sector', 'N/A')}")
        doc.add_paragraph(f"Language: {idea.get('language', 'N/A')}")
        doc.add_heading("AI Recommendations:", level=2)
        doc.add_paragraph(idea.get("recommendations", "No recommendations found."))

        # Save document to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Send file as download
        filename = f"{idea['idea'].replace(' ', '_')}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return "Invalid Idea ID", 404

@app.route("/forget_password", methods=["GET", "POST"])
def forget_password():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        new = request.form.get("new_password", "")
        confirm = request.form.get("confirm_password", "")
        if new != confirm:
            flash("Passwords do not match.", "danger")
            return render_template("forget_password.html")
        users = load_users()
        if username not in users:
            flash("Username not found.", "warning")
            return render_template("forget_password.html")
        users[username]["password"] = generate_password_hash(new)
        save_users(users)
        flash("Password updated successfully! Please login.", "success")
        return redirect(url_for("login"))
    return render_template("forget_password.html")

@app.route("/logout")
def logout():
    session.pop("username", None)
    flash("Logged out successfully.", "success")
    return redirect(url_for("login"))

# -------------------------
# Welcome / Ideas History
# -------------------------
@app.route("/")
def welcome():
    if "username" not in session:
        return redirect(url_for("login"))
    ideas = load_ideas()
    # show latest first
    ideas_sorted = list(reversed(ideas))
    # Note: templates expect idea objects (with attributes like idea.idea or idea['idea'])
    return render_template("welcome.html", username=session["username"], ideas=ideas_sorted)

# -------------------------
# Skills / Jobs
# -------------------------
@app.route("/skills", methods=["GET", "POST"])
def skills():
    if "username" not in session:
        return redirect(url_for("login"))
    if request.method == "POST":
        raw_skill = request.form.get("skill", "").strip().lower()
        if not raw_skill:
            flash("Please enter a valid skill.", "danger")
            return render_template("skills.html", skills_data=skills_data)
        session["skill"] = raw_skill
        companies = skills_data.get(raw_skill, [])
        if not companies:
            flash(f"No jobs available for '{raw_skill}'", "warning")
        return render_template("jobs.html",
                               skill=raw_skill,
                               encoded_skill=encode_skill(raw_skill),
                               companies=companies,
                               skills_data=skills_data)
    return render_template("skills.html", skills_data=skills_data)

@app.route("/upload_project", methods=["GET", "POST"])
def upload_project():
    if "username" not in session:
        return redirect(url_for("login"))

    if request.method == "POST":
        files = request.files.getlist("project_files")
        if not files:
            return "No files uploaded!", 400

        uploads = load_uploads()  # load old JSON
        new_id = len(uploads) + 1
        saved_files = []

        for file in files:
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            saved_files.append(filename)

        # dictionary format
        new_proj = {
            "id": new_id,
            "user": session["username"],
            "files": saved_files,
            "created_at": datetime.utcnow().isoformat() + "Z"
        }

        uploads.append(new_proj)
        save_uploads(uploads) 

        return redirect(url_for("project_history"))

    return render_template("upload_project.html")

@app.route("/project_history")
def project_history():
    uploads = load_uploads()
    history = [
        {
            "id": proj.get("id"),
            "user": proj.get("user"),
            "files": proj.get("files", []),
            "created_at": proj.get("created_at")
        }
        for proj in uploads
    ]
    return render_template("project_history.html", history=history)

@app.route("/download/<filename>")
def download_file(filename):
    if "username" not in session:
        return redirect(url_for("login"))
    user_folder = os.path.join(app.config["UPLOAD_FOLDER"], session["username"])
    return send_file(os.path.join(user_folder, filename), as_attachment=True)

# -------------------------
# Apply (single / multiple) routes
# -------------------------
@app.route("/apply", methods=["GET", "POST"])
def apply_company():
    if "username" not in session:
        return redirect(url_for("login"))
    if request.method == "POST":
        # if coming from company selection (button submit with hidden 'company')
        if "submit" not in request.form and request.form.get("company"):
            company = request.form.get("company")
            return render_template("apply_form.html", company=company, skills_data=skills_data)

        # full form submission
        first_name = request.form.get("first_name", "")
        last_name = request.form.get("last_name", "")
        name = f"{first_name} {last_name}".strip()
        email = request.form.get("email", "")
        country_code = request.form.get("country_code", "")
        phone = request.form.get("phone", "")
        full_phone = f"{country_code}{phone}"
        address = request.form.get("address", "")
        experience = request.form.get("experience", "")
        company = request.form.get("company", "")
        resume = request.files.get("resume")

        resume_filename = ""
        if resume and resume.filename:
            resume_filename = secure_filename(resume.filename)
            resume.save(os.path.join(app.config["UPLOAD_FOLDER"], resume_filename))

        apps_path = os.path.join(DATA_FOLDER, "applications.json")
        try:
            with open(apps_path, "r", encoding="utf-8") as f:
                applications = json.load(f)
        except Exception:
            applications = []

        applications.append({
            "name": name,
            "email": email,
            "phone": full_phone,
            "address": address,
            "experience": experience,
            "company": company,
            "resume": resume_filename
        })

        with open(apps_path, "w", encoding="utf-8") as f:
            json.dump(applications, f, indent=4, ensure_ascii=False)

        session["user_details"] = {"name": name, "email": email, "phone": full_phone}
        return render_template("thankyou.html", name=name, company=company, skills_data=skills_data)

    return render_template("apply_form.html", skills_data=skills_data)

@app.route("/apply_all", methods=["POST"])
def apply_all_or_select():
    if "username" not in session:
        return redirect(url_for("login"))
    skill = request.form.get("skill") or session.get("skill")
    if not skill:
        flash("Please choose a skill first.", "warning")
        return redirect(url_for("skills"))
    return redirect(url_for("select_apply", encoded_skill=encode_skill(skill), all=1))

@app.route("/select_apply/<encoded_skill>", methods=["GET", "POST"])
def select_apply(encoded_skill):
    if "username" not in session:
        return redirect(url_for("login"))
    skill = decode_skill(encoded_skill)
    companies = skills_data.get(skill, [])
    preselect_all = request.args.get("all") == "1"
    user_details = session.get("user_details", {"name": "", "email": "", "phone": ""})
    if not companies:
        flash(f"No companies found for '{skill}'", "warning")
        return redirect(url_for("skills"))
    return render_template("select_apply.html",
                           skill=skill,
                           encoded_skill=encoded_skill,
                           companies=companies,
                           preselect_all=preselect_all,
                           user_details=user_details,
                           skills_data=skills_data)

@app.route("/apply_selected", methods=["POST"])
def apply_selected_route():
    if "username" not in session:
        return redirect(url_for("login"))
    selected_companies = request.form.getlist("selected_companies")
    resume = request.files.get("resume")
    if not selected_companies:
        flash("Select at least one company.", "warning")
        return redirect(url_for("select_apply", encoded_skill=encode_skill(session.get("skill", ""))))
    if not resume or resume.filename == "":
        flash("Please upload resume.", "warning")
        return redirect(url_for("select_apply", encoded_skill=encode_skill(session.get("skill", ""))))

    resume_filename = secure_filename(resume.filename)
    resume.save(os.path.join(app.config["UPLOAD_FOLDER"], resume_filename))

    apps_path = os.path.join(DATA_FOLDER, "applications.json")
    try:
        with open(apps_path, "r", encoding="utf-8") as f:
            apps = json.load(f)
    except Exception:
        apps = []

    for comp in selected_companies:
        apps.append({
            "name": session.get("user_details", {}).get("name", ""),
            "email": session.get("user_details", {}).get("email", ""),
            "phone": session.get("user_details", {}).get("phone", ""),
            "address": "",
            "experience": "",
            "company": comp,
            "resume": resume_filename
        })
    with open(apps_path, "w", encoding="utf-8") as f:
        json.dump(apps, f, indent=4, ensure_ascii=False)

    return render_template("thankyou.html",
                           name=session.get("user_details", {}).get("name", "User"),
                           company=", ".join(selected_companies),
                           skills_data=skills_data)

# -------------------------
# Resume check endpoint (used by skills.html)
# -------------------------
@app.route("/check_resume", methods=["POST"])
def check_resume():
    if "resume" not in request.files:
        return jsonify({"error": "No resume uploaded"}), 400
    file = request.files["resume"]
    filename = secure_filename(file.filename)
    if filename == "":
        return jsonify({"error": "No selected file"}), 400
    save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(save_path)
    # Optionally extract text and evaluate; simplified here
    return jsonify({"message": "Resume uploaded successfully!", "filename": filename})

# -------------------------
# AI Recommendations (Gemini - optional)
# -------------------------
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", None)

@app.route("/recommend", methods=["POST"])
def recommend():
    if "username" not in session:
        return jsonify({"error": "login required"}), 401
    data = request.get_json(silent=True) or request.form or {}
    idea_text = (data.get("idea") or "").strip()
    if not idea_text:
        return jsonify({"recommendations": "Please provide an idea."}), 400

    # ---- 1. generate recommendations ----
    if not GEMINI_API_KEY:
        recommendations = f"AI key not configured. Example recommendations for: {idea_text}\n\n1) Define scope.\n2) Choose tech stack.\n3) Build MVP."
    else:
        payload = {
            "contents": [{
                "parts": [{"text": f"User idea: {idea_text}\nGenerate structured recommendations, improvements, and next steps."}]
            }]
        }
        headers = {"Content-Type": "application/json"}
        try:
            resp = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}",
                headers=headers,
                json=payload,
                timeout=30
            )
            resp.raise_for_status()
            result = resp.json()
            candidates = result.get("candidates") or []
            recommendations = ""
            if candidates:
                parts = candidates[0].get("content", {}).get("parts", [])
                if parts:
                    recommendations = parts[0].get("text", "")
            if not recommendations:
                recommendations = json.dumps(result, indent=2)[:2000]
        except Exception as e:
            return jsonify({"recommendations": f"Error: {str(e)}"}), 502

    # ---- 2. save idea immediately ----
    ideas = load_ideas()
    new_id = next_idea_id(ideas)
    new_item = {
        "id": new_id,
        "user": session["username"],
        "idea": idea_text,
        "sector": "",
        "language": "English",
        "recommendations": recommendations,
        "created_at": datetime.utcnow().isoformat() + "Z"
    }
    ideas.append(new_item)
    save_ideas(ideas)

    # ---- 3. return response ----
    return jsonify({"recommendations": recommendations, "idea_id": new_id})

# -------------------------
# Ideas page (enter idea -> call /recommend via JS)
# -------------------------
@app.route("/ideas", methods=["GET"])
def ideas():
    if "username" not in session:
        return redirect(url_for("login"))
    return render_template("recommend.html", skills_data=skills_data)

# collaboration_form (GET) - show collaboration form, accepts idea + recommendations via query params
@app.route("/collaboration_form", methods=["GET"])
def collaboration_form():
    if "username" not in session:
        return redirect(url_for("login"))
    idea = request.args.get("idea", "")
    recommendations = request.args.get("recommendations", "")
    idea_id = request.args.get("idea_id", "")
    return render_template(
        "collaboration.html",
        idea=idea,
        recommendations=recommendations,
        idea_id=idea_id
    )
    
@app.route("/clear_history", methods=["POST"])
def clear_history():
    if "username" not in session:
        return redirect(url_for("login"))
    
    # Empty list → overwrite ideas.json
    save_ideas([])
    return redirect(url_for("welcome"))
    
# collaboration (POST) - save to ideas.json with id and recommendations
@app.route("/collaboration", methods=["POST"])
def collaboration():
    if "username" not in session:
        return redirect(url_for("login"))

    data = request.form
    idea_id = data.get("idea_id")
    sector = data.get("sector", "")
    language = data.get("language", "English")

    ideas = load_ideas()

    # ✅ if idea_id exists → update the existing idea instead of adding duplicate
    if idea_id:
        for idea in ideas:
            if str(idea.get("id")) == str(idea_id):
                idea["sector"] = sector
                idea["language"] = language
                idea["updated_at"] = datetime.utcnow().isoformat() + "Z"
                break
    else:
        # (backup: only if somehow no idea_id came)
        new_id = next_idea_id(ideas)
        new_item = {
            "id": new_id,
            "user": session["username"],
            "idea": data.get("idea", ""),
            "sector": sector,
            "language": language,
            "recommendations": data.get("recommendations", ""),
            "created_at": datetime.utcnow().isoformat() + "Z"
        }
        ideas.append(new_item)

    save_ideas(ideas)
    return redirect(url_for("welcome"))

# -------------------------
# Start project: downloads a .docx containing the idea + recommendations
# Template & welcome expect to call this with idea_id
# -------------------------
@app.route("/start_project/<int:idea_id>")
def start_project(idea_id):
    ideas = load_ideas()
    # find idea by numeric id
    idea = next((i for i in ideas if int(i.get("id", 0)) == int(idea_id)), None)
    if not idea:
        flash("Idea not found.", "danger")
        return redirect(url_for("welcome"))

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading(f"Project: {idea.get('idea', '')}", level=0)
    doc.add_paragraph(f"Submitted by: {idea.get('user', '')}")
    doc.add_paragraph(f"Sector: {idea.get('sector', '')}")
    doc.add_paragraph(f"Preferred Language: {idea.get('language', '')}")
    if idea.get("created_at"):
        doc.add_paragraph(f"Created At: {idea.get('created_at')}")
    doc.add_heading("AI Recommendations", level=1)
    recs = idea.get("recommendations", "")
    if recs:
        for line in recs.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("No recommendations saved for this idea.")

    doc.add_heading("Next Steps (Suggested)", level=1)
    doc.add_paragraph("1) Create a repo/folder locally (e.g., in VSCode).")
    doc.add_paragraph("2) Copy these recommendations into README.md or project plan.")
    doc.add_paragraph("3) Start implementing modules one-by-one and commit often.")
    doc.add_paragraph("4) Iterate with the AI to refine next steps as you progress.")

    doc.save(buf)
    buf.seek(0)

    filename = _safe_filename(idea.get("idea", "project")) + ".docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# -------------------------
# (Optional) static feed page to view community feed (if you have show_feed)
# -------------------------
@app.route("/feed")
def feed():
    if "username" not in session:
        return redirect(url_for("login"))
    try:
        feed_data = show_feed()
    except Exception:
        feed_data = []
    return render_template("feed.html", feed=feed_data)

# -------------------------
# Start the app
# -------------------------
if __name__ == "__main__":
    # create empty ideas file if missing (safe initialization)
    if not os.path.exists(IDEAS_FILE):
        save_ideas([])
    app.run(debug=True)
